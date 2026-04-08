"""
Claude AI integration — document classification + streaming query.
"""
import os
import io
import json
import base64
from pathlib import Path
from typing import Any, AsyncIterator, Optional
from dotenv import load_dotenv
import anthropic

# Load .env from the same directory as this file — works regardless of CWD
load_dotenv(Path(__file__).parent / ".env")

CLASSIFY_MODEL = "claude-haiku-4-5"   # Haiku for fast classification


def _extract_text(filename: str, raw_bytes: bytes) -> str:
    """Extract readable text from various file formats."""
    ext = Path(filename).suffix.lower()

    # Word (.docx)
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n".join(paragraphs) or "(empty document)"
        except Exception as e:
            return f"(Could not read Word file: {e})"

    # Excel (.xlsx)
    if ext in (".xlsx", ".xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(c) for c in row if c is not None)
                    if row_text.strip():
                        lines.append(row_text)
            return "\n".join(lines) or "(empty spreadsheet)"
        except Exception as e:
            return f"(Could not read Excel file: {e})"

    # CSV / plain text
    if ext in (".csv", ".txt", ".md", ".json", ".xml", ".html"):
        try:
            return raw_bytes.decode("utf-8", errors="replace")
        except Exception:
            return "(Could not decode file)"

    # Fallback — try plain text decode
    try:
        decoded = raw_bytes.decode("utf-8", errors="replace")
        # If it looks like binary garbage, say so
        printable = sum(1 for c in decoded if c.isprintable() or c in "\n\r\t")
        if printable / max(len(decoded), 1) < 0.7:
            return f"(Binary file — classify by filename: {filename})"
        return decoded
    except Exception:
        return f"(Binary file — classify by filename: {filename})"
QUERY_MODEL   = "claude-haiku-4-5"   # Haiku for fast conversational Q&A

_anthropic: Optional[anthropic.AsyncAnthropic] = None


def get_anthropic() -> anthropic.AsyncAnthropic:
    global _anthropic
    if _anthropic is None:
        _anthropic = anthropic.AsyncAnthropic(
            api_key=os.getenv("ANTHROPIC_API_KEY")
        )
    return _anthropic


# ─── Classification prompt ────────────────────────────────────────────────────

CLASSIFY_SYSTEM = """You are an expert document analyst. Analyze this document and return \
ONLY a JSON object with these exact fields:
- category: broad category (Finance, Healthcare, Legal, Retail, HR, Education, \
Real Estate, Insurance, Technology, Personal, Other)
- document_type: specific type (e.g. 'Vendor Invoice', 'Medical Report', \
'Employment Contract', 'Bank Statement', 'Resume', 'Receipt')
- confidence: float 0-1
- summary: one sentence describing this specific document
- extracted_fields: object with ALL relevant fields you can find \
(dates, amounts, names, IDs, addresses, line items — everything)
- suggested_questions: array of exactly 5 questions a user would actually want \
to ask about THIS specific document (make them concrete and useful)
- theme_color: hex color that fits the category \
(Finance=#3b82f6, Healthcare=#10b981, Legal=#8b5cf6, \
Retail=#f59e0b, HR=#ec4899, Education=#14b8a6, \
Real Estate=#f97316, Insurance=#64748b, Technology=#6366f1, Personal=#a855f7, Other=#6366f1)

Return ONLY the JSON object. No markdown, no code fences, no explanation."""


async def classify_document(
    file_b64: str,
    media_type: str,
    filename: str,
) -> dict[str, Any]:
    """
    Send file to Claude for classification + field extraction in one call.
    Returns the parsed JSON response dict, including 'full_text' for text-based files.
    """
    client = get_anthropic()

    # Claude vision only accepts these four image types
    VISION_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp"}

    full_text: str = ""  # will hold extracted text for non-image/pdf files

    # Build the message content
    if media_type in VISION_TYPES:
        # Native image vision
        content = [
            {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": file_b64,
                },
            },
            {
                "type": "text",
                "text": f"Filename: {filename}\n\nAnalyze this document and return the JSON as instructed.",
            },
        ]
    elif media_type == "application/pdf":
        # PDF document type (supported by Claude 3.5+)
        content = [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": file_b64,
                },
            },
            {
                "type": "text",
                "text": f"Filename: {filename}\n\nAnalyze this document and return the JSON as instructed.",
            },
        ]
    else:
        # Extract readable text based on file type
        raw_bytes = base64.b64decode(file_b64)
        full_text = _extract_text(filename, raw_bytes)
        content = [
            {
                "type": "text",
                "text": f"Filename: {filename}\n\nDocument content:\n{full_text[:12000]}\n\nAnalyze this document and return the JSON as instructed.",
            }
        ]

    message = await client.messages.create(
        model=CLASSIFY_MODEL,
        max_tokens=4096,
        system=CLASSIFY_SYSTEM,
        messages=[{"role": "user", "content": content}],
    )

    raw = message.content[0].text.strip()

    # Strip markdown code fences if Claude adds them despite the prompt
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    result = json.loads(raw)

    # Attach full text so it can be stored and used during Q&A
    if full_text:
        result["full_text"] = full_text[:20000]  # cap at 20k chars

    return result


# ─── Streaming query ──────────────────────────────────────────────────────────

QUERY_SYSTEM_TEMPLATE = """You've read this document thoroughly:

Type: {document_type} | Category: {category}
Summary: {summary}

Extracted Fields:
{extracted_fields}
{full_text_section}
You are a knowledgeable friend having a real conversation — not an AI giving a presentation. \
Think of how a doctor friend explains something at dinner, not how a medical textbook reads.

How to respond:
- Answer only what was asked. Don't volunteer extra information.
- 2-3 sentences maximum. If the answer is one sentence, that's perfect.
- Speak naturally, like you're texting a smart friend. Contractions are fine.
- Use one specific number or fact from the data if it's relevant — not all of them.
- Never start with "So", "Great question", "Certainly", or restate the question.
- If the answer is simple, keep it simple. Don't pad it out.
- No lists, no bullet points, no headers. Just talk.
- If you don't know, say "That's not in the document" and stop there."""


async def stream_query(
    question: str,
    document_context: dict[str, Any],
    conversation_history: Optional[list[dict]] = None,
) -> AsyncIterator[str]:
    """
    Stream a Claude response token-by-token.
    Yields text chunks as they arrive from the API.
    """
    client = get_anthropic()

    doc_type = document_context.get("document_type", "Document")
    category = document_context.get("category", "General")
    summary = document_context.get("summary", "")
    extracted = document_context.get("extracted_fields", {})
    full_text = document_context.get("full_text", "")

    extracted_str = "\n".join(
        f"  {k}: {v}" for k, v in extracted.items()
    ) if extracted else "  (no fields extracted)"

    # Include full document text for text-based files (docx, xlsx, csv, txt, etc.)
    full_text_section = (
        f"\nFull Document Text:\n{full_text[:15000]}\n"
        if full_text else ""
    )

    system_prompt = QUERY_SYSTEM_TEMPLATE.format(
        document_type=doc_type,
        category=category,
        summary=summary,
        extracted_fields=extracted_str,
        full_text_section=full_text_section,
    )

    messages = []
    if conversation_history:
        messages.extend(conversation_history)
    messages.append({"role": "user", "content": question})

    async with client.messages.stream(
        model=QUERY_MODEL,
        max_tokens=300,
        system=system_prompt,
        messages=messages,
    ) as stream:
        async for text in stream.text_stream:
            yield text
